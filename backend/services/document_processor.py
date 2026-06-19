import os
import shutil
from pathlib import Path
from typing import BinaryIO
import PyPDF2
from docx import Document as DocxDocument
import nltk
from nltk.tokenize import sent_tokenize

class DocumentProcessor:
    """Handles document file processing and text extraction"""
    
    def __init__(self, upload_dir: Path):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    def save_uploaded_file(self, file: BinaryIO, filename: str) -> str:
        """Save uploaded file to disk"""
        file_path = self.upload_dir / filename
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        return str(file_path)
    
    def extract_text_from_pdf(self, file_path: str) -> tuple[str, int]:
        """Extract text and page count from PDF"""
        text = ""
        page_count = 0
        
        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                page_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        
        return text, page_count
    
    def extract_text_from_docx(self, file_path: str) -> tuple[str, None]:
        """Extract text from DOCX"""
        text = ""
        
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        
        return text, None
    
    def extract_text_from_txt(self, file_path: str) -> tuple[str, None]:
        """Extract text from TXT"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            return text, None
        except Exception as e:
            print(f"Error reading TXT: {e}")
            return "", None
    
    def extract_text(self, file_path: str, file_type: str) -> tuple[str, int]:
        """Main method to extract text based on file type"""
        file_type = file_type.lower().lstrip(".")
        
        if file_type == "pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == "docx":
            text, _ = self.extract_text_from_docx(file_path)
            return text, None
        elif file_type == "txt":
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def get_word_count(text: str) -> int:
        """Count words in text"""
        return len(text.split())
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Basic text cleaning"""
        # Remove extra whitespace
        text = " ".join(text.split())
        # Remove special characters but keep basic punctuation
        return text
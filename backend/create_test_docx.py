from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a Document
doc = Document()

# Add title
title = doc.add_heading('Artificial Intelligence in 2024', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add paragraphs
p1 = doc.add_paragraph(
    "Artificial Intelligence (AI) has become one of the most transformative "
    "technologies of our time. From healthcare to finance, AI is revolutionizing "
    "how we work and live."
)

doc.add_heading('Key Developments in AI', level=1)

doc.add_paragraph(
    "Large Language Models - GPT models have demonstrated remarkable capabilities "
    "in understanding and generating human language.",
    style='List Bullet'
)

doc.add_paragraph(
    "Computer Vision - AI systems can now interpret images and videos with "
    "superhuman accuracy.",
    style='List Bullet'
)

doc.add_paragraph(
    "Autonomous Systems - Self-driving cars and robots are becoming more sophisticated.",
    style='List Bullet'
)

doc.add_heading('Impact on Different Industries', level=1)

doc.add_paragraph(
    "Healthcare: AI is assisting in diagnosis, drug discovery, and personalized medicine."
)

doc.add_paragraph(
    "Finance: Algorithmic trading, fraud detection, and risk assessment."
)

doc.add_paragraph(
    "Education: Personalized learning, automated grading, and intelligent tutoring systems."
)

doc.add_paragraph(
    "Manufacturing: Predictive maintenance, quality control, and optimization."
)

doc.add_heading('Challenges and Ethical Concerns', level=1)

doc.add_paragraph(
    "While AI offers tremendous potential, it also raises important questions about "
    "privacy, bias, job displacement, and responsible AI development."
)

doc.add_paragraph(
    "Organizations must consider ethical implications and work towards creating "
    "AI systems that are fair, transparent, and beneficial to society."
)

# Save document
docx_path = "sample_document.docx"
doc.save(docx_path)
print(f"✓ Created: {docx_path}")